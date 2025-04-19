from docx import Document
import os
import uuid
import json

def save_response_to_word(genai_response, filename=None):
    
    if genai_response.startswith('"') and genai_response.endswith('"'):
        genai_response = genai_response[1:-1]
        genai_response = genai_response.replace('\\"', '"')

    document = Document()
    document.add_heading("Business Document Analysis", 0)

    if isinstance(genai_response, str):
        try:
            if not genai_response.strip():
                raise ValueError("The genai_response is empty or invalid.")
            genai_response = json.loads(genai_response)
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON data: {e}")

    # modules section
    modules = genai_response.get("modules", [])
    if modules:
        document.add_heading("Modules", level=1)
        for module in modules:
            document.add_heading(module["name"], level=2)
            document.add_paragraph(module["functionality"])
    
    # use cases section
    use_cases = genai_response.get("use_cases", [])
    if use_cases:
        document.add_heading("Use Cases", level=1)
        for use_case in use_cases:
            document.add_paragraph(f"• {use_case}")
    
    # stake holders section
    stakeholders = genai_response.get("stakeholders", [])
    if stakeholders:
        document.add_heading("Stakeholders", level=1)
        for stakeholder in stakeholders:
            document.add_heading(stakeholder["role"], level=2)
            document.add_paragraph(stakeholder["responsibility"])
    
    # risks section
    risks = genai_response.get("risks", [])
    if risks:
        document.add_heading("Risks Involved", level=1)
        for risk in risks:
            document.add_paragraph(f"• {risk}")
    
    # priority section
    timeline_priority = genai_response.get("timeline_priority", [])
    if timeline_priority:  
        document.add_heading("Timeline and Priority", level=1)
        print("heading crossed")
        for mod in timeline_priority:
            document.add_heading(mod["module"], level=2)
            document.add_heading(mod["priority"], level=3)
            document.add_paragraph(mod["expected_timeline"])
    
    # Save document
    if not filename:
        filename = f"{uuid.uuid4().hex}.docx"
    file_path = os.path.join("media", filename)
    document.save(file_path)
    return file_path
